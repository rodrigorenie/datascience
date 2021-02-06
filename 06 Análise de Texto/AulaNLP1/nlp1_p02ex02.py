import docx
import nltk
from nltk import word_tokenize

noticia1 = docx.Document('../Dados/Noticia_1.docx')
noticia2 = docx.Document('../Dados/Noticia_2.docx')

noticia1_text = '\n'.join([p.text for p in noticia1.paragraphs])
noticia2_text = '\n'.join([p.text for p in noticia2.paragraphs])

noticia1_tokens = word_tokenize(noticia1_text)
noticia2_tokens = word_tokenize(noticia1_text)

for w, t in nltk.pos_tag(noticia1_tokens):
    print('{:15s} {}'.format(w, t))