import docx
import nltk
from nltk import word_tokenize


file = '../Dados/Textos para copiar.docx'
doc = docx.Document(file)

text_e = doc.paragraphs[4].text
text_p = doc.paragraphs[6].text
tokens_p = word_tokenize(text_p)
tokens_e = word_tokenize(text_e)

for w, t in nltk.pos_tag(tokens_e):
    print('{:15s} {}'.format(w, t))

nltk.help.upenn_tagset('NNPS')
