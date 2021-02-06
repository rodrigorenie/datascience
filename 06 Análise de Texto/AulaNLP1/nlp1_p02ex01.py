from nltk.corpus import machado
from nltk import PorterStemmer, LancasterStemmer, RSLPStemmer
from nltk import word_tokenize
import docx

#
# Carrega os parágrafos de Memórias Póstumas de Brás Cubas
#

book_fileid = 'romance/marm05.txt'

# Carrega todos os parágrafos do livro, já tokenizados
book_paras = machado.paras(book_fileid)

# A posição 17 é o primeiro parágrafo do primeiro capítulo, portanto:
tokens = book_paras[17][0] + book_paras[18][0]

#
# Carrega os parágrafos de Notícia 2
#
noticia2 = '../Dados/Noticia_2.docx'
doc = docx.Document(noticia2)
text = doc.paragraphs[2].text + doc.paragraphs[3].text
tokens2 = word_tokenize(text)

#
# Executa os Stemmers
#
porter = PorterStemmer()
stems_porter = [porter.stem(t) for t in tokens]

lancaster = LancasterStemmer()
stems_lancaster = [lancaster.stem(t) for t in tokens]

rslp = RSLPStemmer()
stems_rslp = [rslp.stem(t) for t in tokens]

print('{:18s} {:18s} {:18s} {}'.format('Tokens', 'Porter', 'Lancaster', 'RSLP'))
for t, p, l, r in zip(tokens, stems_porter, stems_lancaster, stems_rslp):
    print('{:18s} {:18s} {:18s} {}'.format(t, p, l, r))
