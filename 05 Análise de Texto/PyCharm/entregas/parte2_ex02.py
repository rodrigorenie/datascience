# Rodrigo Renie de Braga Pinto
# TEXT ANALYSIS(Apostila)Parte 2.docx
# Exercitando 2
# Execute o que se pede.
# Utilize o arquivo Noticia_1 disponível na pasta de dados da turma e realize
# as tarefas de POS e NER.
# Salve o resultado de POS em um arquivo, com <seu_nome>_POS_Noticia1
# Salve o resultado de NER em um arquivo, com <seu_nome>_NER_Noticia1

import docx
import nltk

noticia1 = docx.Document('dados/Noticia_1.docx')

noticia1_tokens = '\n'.join([p.text for p in noticia1.paragraphs])
noticia1_tokens = nltk.sent_tokenize(noticia1_tokens)
noticia1_tokens = [nltk.word_tokenize(t) for t in noticia1_tokens]
print(noticia1_tokens)

noticia1_pos = [nltk.pos_tag(s) for s in noticia1_tokens]
noticia1_ner = [nltk.ne_chunk(s) for s in noticia1_pos]
print(noticia1_pos)

exit()
posdoc = docx.Document()
for pos in noticia1_pos:
    posdoc.add_paragraph(str(pos))
posdoc.save('dados/RodrigoRenieDeBragaPinto_POS_Noticia1.docx')

nerdoc = docx.Document()
for ner in noticia1_ner:
    nerdoc.add_paragraph(str(ner))
nerdoc.save('dados/RodrigoRenieDeBragaPinto_NER_Noticia1.docx')

